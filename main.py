from docx import Document

# Create a new Word document
doc = Document()

# Add Title
doc.add_heading('Container Orchestration Report', 0)

# Table of Contents
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph("""
1. Introduction
   - 1.1 Concept of Microservices and Containerization
   - 1.2 Importance of Container Orchestration
   - 1.3 Logging and Monitoring of Services

2. Docker Compose Deployment
   - 2.1 Overview of Docker Compose
   - 2.2 Step-by-Step Deployment Guide
     - 2.2.1 Dockerfile Creation
     - 2.2.2 Launching Containers (Flightservice and Database)
     - 2.2.3 Database Setup (flight.sql)
     - 2.2.4 Testing Using Postman
     - 2.2.5 Pushing Flightservice Image to DockerHub
     - 2.2.6 Managing Services with Docker Compose
     - 2.2.7 Integrating ELK Stack for Logging and Monitoring

3. Docker Swarm Deployment
   - 3.1 Overview of Docker Swarm and Core Components
   - 3.2 Step-by-Step Deployment Guide
     - 3.2.1 Deploying Microservice with Docker Swarm
     - 3.2.2 Demonstrating Scaling, Upgrade, and Rollback
     - 3.2.3 Commands and Screenshots

4. Kubernetes Deployment
   - 4.1 Overview of Kubernetes and Core Components
   - 4.2 Step-by-Step Deployment Guide
     - 4.2.1 Deploying Microservice with Kubernetes
     - 4.2.2 Demonstrating Scaling, Upgrade, and Rollback
     - 4.2.3 Using Helm Charts for Simplified Deployment
   - 4.3 YAML Configuration Files and Command-Line Examples

5. Evaluation & Conclusion
   - 5.1 Task Completion Analysis
   - 5.2 Problems Encountered
   - 5.3 Comparison of Kubernetes and Docker Swarm
     - 5.3.1 Ease of Use
     - 5.3.2 Learning Curve
     - 5.3.3 Features

6. References
""")

# Add Section Placeholders
sections = [
    ("1. Introduction", [
        "1.1 Concept of Microservices and Containerization",
        "1.2 Importance of Container Orchestration",
        "1.3 Logging and Monitoring of Services"
    ]),
    ("2. Docker Compose Deployment", [
        "2.1 Overview of Docker Compose",
        "2.2 Step-by-Step Deployment Guide",
        "2.2.1 Dockerfile Creation",
        "2.2.2 Launching Containers (Flightservice and Database)",
        "2.2.3 Database Setup (flight.sql)",
        "2.2.4 Testing Using Postman",
        "2.2.5 Pushing Flightservice Image to DockerHub",
        "2.2.6 Managing Services with Docker Compose",
        "2.2.7 Integrating ELK Stack for Logging and Monitoring"
    ]),
    ("3. Docker Swarm Deployment", [
        "3.1 Overview of Docker Swarm and Core Components",
        "3.2 Step-by-Step Deployment Guide",
        "3.2.1 Deploying Microservice with Docker Swarm",
        "3.2.2 Demonstrating Scaling, Upgrade, and Rollback",
        "3.2.3 Commands and Screenshots"
    ]),
    ("4. Kubernetes Deployment", [
        "4.1 Overview of Kubernetes and Core Components",
        "4.2 Step-by-Step Deployment Guide",
        "4.2.1 Deploying Microservice with Kubernetes",
        "4.2.2 Demonstrating Scaling, Upgrade, and Rollback",
        "4.2.3 Using Helm Charts for Simplified Deployment",
        "4.3 YAML Configuration Files and Command-Line Examples"
    ]),
    ("5. Evaluation & Conclusion", [
        "5.1 Task Completion Analysis",
        "5.2 Problems Encountered",
        "5.3 Comparison of Kubernetes and Docker Swarm",
        "5.3.1 Ease of Use",
        "5.3.2 Learning Curve",
        "5.3.3 Features"
    ]),
    ("6. References", [])
]

for section_title, subsections in sections:
    doc.add_heading(section_title, level=1)
    for subsection in subsections:
        doc.add_heading(subsection, level=2)
        doc.add_paragraph("Placeholder for content.\n")

# Save the document
file_path = "/mnt/data/container_orchestration_report_template.docx"
doc.save(file_path)

# file_path
